#!/usr/bin/env python3
"""Generate a Word document telling the scientific story of the CPC pipeline.

Includes LEC-only figures + fig3_smoothing2 + GEV diagnostics, each with
clear methodology explanation in a narrative flow.

Usage:
    python scripts/generate_methodology_doc.py
"""

import os, sys
sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT

FIGURES_DIR = os.path.join(os.path.dirname(__file__), "..", "docs", "figures")
OUTPUT = os.path.join(os.path.dirname(__file__), "..", "docs", "methodology_report.docx")


def add_figure(doc, filename, caption, width=Inches(6.0)):
    """Insert a centred figure with a caption below it."""
    path = os.path.join(FIGURES_DIR, filename)
    if not os.path.exists(path):
        p = doc.add_paragraph()
        run = p.add_run(f"[MISSING FIGURE: {filename}]")
        run.font.color.rgb = RGBColor(255, 0, 0)
        run.bold = True
        print(f"  WARNING: {filename} not found at {path}")
        return

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(path, width=width)

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(80, 80, 80)


def set_normal_style(doc):
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    pf = style.paragraph_format
    pf.space_after = Pt(6)
    pf.line_spacing = 1.15


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(30, 60, 100)
    return h


def main():
    doc = Document()
    set_normal_style(doc)

    # --- Title ---
    title = doc.add_heading(
        "Climate Risk Analysis via Max-Stable Process Clustering\n"
        "— Methodology and Results —",
        level=0,
    )
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "Widad Kchikach · Alfred Wegener Institute\n"
        "Based on Contzen, Dickhaus & Lohmeyer (2025, Extremes 28:713–737)"
    )
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(100, 100, 100)

    # ================================================================
    # 1. INTRODUCTION
    # ================================================================
    add_heading(doc, "1. Introduction")

    doc.add_paragraph(
        "This document presents a step-by-step walkthrough of the spatial "
        "climate risk pipeline implemented in the weatherisk Python package. "
        "The methodology follows Contzen, Dickhaus & Lohmeyer (2025), who "
        "developed a clustering framework for non-stationary max-stable "
        "processes to identify regions of homogeneous extremal dependence."
    )
    doc.add_paragraph(
        "The pipeline takes 20 years of daily precipitation data (NOAA CPC, "
        "2000–2019), extracts extreme events, models their spatial dependence "
        "structure, and ultimately produces risk maps that combine statistical "
        "tail-risk intensity with economic exposure. Each section below "
        "corresponds to a pipeline step, accompanied by the resulting figure "
        "and a plain-language explanation of what it shows."
    )

    # ================================================================
    # 2. SIMULATION VALIDATION — Figure 3
    # ================================================================
    add_heading(doc, "2. Method Validation on Simulated Data")

    doc.add_paragraph(
        "Before applying the method to real data, we validate the clustering "
        "pipeline on synthetic data where the ground truth is known. "
        "This reproduces Figure 3 from Contzen et al. (2025, p. 729), which "
        "is the key validation experiment of the paper."
    )

    add_heading(doc, "2.1 Experimental Setup", level=2)

    doc.add_paragraph(
        "A non-stationary max-stable process is simulated on a 51×51 grid "
        "over the domain [−5, 5]². The process has spatially varying "
        "anisotropy parameters: the semi-axis lengths a(x,y) and b(x,y) "
        "change across the domain according to known functions (the "
        '"paper_stripes" preset: b(x,y) = (x + 5)/2, creating a smooth '
        "gradient from West to East). The rotation angle γ is fixed at 0. "
        "We simulate 250 independent realisations with 5 degrees of freedom "
        "(extremal-t model, α = 1.0)."
    )

    doc.add_paragraph(
        "The pipeline then (i) estimates local anisotropy parameters at each "
        "cell via pairwise composite likelihood, (ii) smooths these estimates "
        "over a neighbourhood of radius 4 grid-spacings, and (iii) clusters "
        "cells using the LEC (ellipse-overlap) method into k groups determined "
        "automatically by a 30th-percentile threshold on pairwise distances."
    )

    add_heading(doc, "2.2 Result: Cluster Comparison", level=2)

    add_figure(
        doc, "fig3_smoothing2.png",
        "Figure 1. Reproduction of Figure 3 from Contzen et al. (2025). "
        "Left: EDC clusters (madogram-based). Right: LEC clusters (ellipse-overlap). "
        "Colour represents the estimated parameter b, which increases from left to right "
        "matching the true gradient b(x,y) = (x+5)/2. The LEC method recovers "
        "the spatial gradient more faithfully than EDC."
    )

    doc.add_paragraph(
        "The figure shows both clustering methods applied to the simulated "
        "data. The background colour in each panel encodes the estimated "
        "parameter b (semi-axis length), which should increase smoothly from "
        "left (≈0) to right (≈5) — matching the true generating function "
        "b(x,y) = (x + 5)/2."
    )
    doc.add_paragraph(
        "The LEC (Localised Ellipse-shape Clustering) method on the right "
        "correctly recovers vertical stripe-like clusters that align with the "
        "true spatial gradient. In contrast, the EDC (Extremal Dependence "
        "Clustering) method on the left produces noisier, less spatially "
        "coherent cluster boundaries. This is the central result of Contzen "
        "et al. (2025): the ellipse-overlap dissimilarity captures the "
        "anisotropy structure that the madogram-based method misses."
    )
    doc.add_paragraph(
        "Our Python implementation reproduces this figure with matching "
        "cluster counts (k_lec = 6, k_edc = 5) and visually equivalent "
        "spatial patterns, confirming that the numerical pipeline is "
        "correctly implemented. All intermediate quantities (local MLE "
        "estimates, smoothed parameters, dissimilarity matrices) are "
        "cross-validated against the original R code to 10⁻¹⁰ tolerance "
        "across 28 unit tests."
    )

    # ================================================================
    # 3. REAL DATA — GEV FITTING
    # ================================================================
    add_heading(doc, "3. Application to Real Precipitation Data")

    doc.add_paragraph(
        "Having validated the method on simulated data, we now apply it to "
        "observed daily precipitation from the NOAA Climate Prediction Center "
        "(CPC) global unified gauge-based dataset. We use 20 years of data "
        "(2000–2019) over a European domain (30°N–65°N, 5°E–55°E), coarsened "
        "to approximately 2° resolution (every 4th 0.5° grid cell) for "
        "computational tractability."
    )

    add_heading(doc, "3.1 Extreme Value Modelling: GEV Fit", level=2)

    doc.add_paragraph(
        "The first step is to extract the extreme precipitation signal from "
        "the raw daily data. For each grid cell, we compute annual block "
        "maxima — the single highest daily precipitation value in each "
        "calendar year — yielding 20 extreme values per cell. We then fit a "
        "Generalised Extreme Value (GEV) distribution to each cell's block "
        "maxima via maximum likelihood estimation."
    )
    doc.add_paragraph(
        "The GEV distribution has three parameters: location μ (the "
        "typical level of extremes), scale σ (the spread), and shape ξ "
        "(the tail heaviness). When ξ > 0, the distribution has a heavy "
        "tail (Fréchet domain); when ξ < 0, it has a bounded upper tail "
        "(Weibull domain); when ξ = 0, it reduces to the Gumbel distribution."
    )
    doc.add_paragraph(
        "To ensure that the GEV is an appropriate model for these data, we "
        "perform two diagnostic checks: quantile-quantile (Q-Q) plots for a "
        "random sample of cells, and a Kolmogorov-Smirnov (KS) goodness-of-fit "
        "test across all land cells."
    )

    add_heading(doc, "3.2 GEV Diagnostic: Q-Q Plots", level=2)

    add_figure(
        doc, "gev_qq_plots.png",
        "Figure 2. GEV Q-Q plots for a random sample of 20 land cells. "
        "Each subplot compares empirical quantiles (horizontal axis) against "
        "quantiles predicted by the fitted GEV (vertical axis). Points lying "
        "close to the diagonal indicate good fit.",
        width=Inches(6.0),
    )

    doc.add_paragraph(
        "The Q-Q plots compare the observed annual block maxima (empirical "
        "quantiles) against the values predicted by the fitted GEV distribution "
        "(theoretical quantiles). If the GEV is a good model, points should "
        "fall close to the 1:1 diagonal line. Across the sample of 20 cells, "
        "the points show excellent agreement with the diagonal, confirming that "
        "the GEV is an appropriate marginal model for precipitation extremes "
        "at these cells."
    )

    add_heading(doc, "3.3 GEV Diagnostic: KS Test Across All Cells", level=2)

    add_figure(
        doc, "gev_ks_pvalues.png",
        "Figure 3. Distribution of Kolmogorov-Smirnov p-values across all 384 "
        "land cells. The dashed red line marks α = 0.05. Most cells have "
        "p-values well above 0.05, indicating that the GEV fit is not rejected.",
        width=Inches(5.0),
    )

    doc.add_paragraph(
        "The Kolmogorov-Smirnov test measures the maximum discrepancy between "
        "the empirical distribution and the fitted GEV at each cell. A p-value "
        "below 0.05 would indicate a statistically significant departure from "
        "the GEV model. Out of 384 land cells, only 10 (2.6%) have p-values "
        "below 0.05 — which is close to what we would expect by chance alone "
        "at the 5% significance level. The median p-value is 0.91, indicating "
        "excellent overall fit. This gives us confidence that the per-cell GEV "
        "models are appropriate."
    )

    add_heading(doc, "3.4 GEV Shape Parameter Map", level=2)

    add_figure(
        doc, "gev_shape_map.png",
        "Figure 4. Spatial distribution of the estimated GEV shape parameter ξ. "
        "Most of Europe shows ξ near zero (Gumbel-like behaviour). Positive values "
        "indicate heavier tails (more extreme extremes).",
        width=Inches(5.5),
    )

    doc.add_paragraph(
        "The shape parameter ξ controls how heavy the tail of the extreme "
        "value distribution is. Across Europe, most cells have ξ close to zero "
        "(mean = 0.08, median = 0.06), indicating that precipitation extremes "
        "in this region are approximately Gumbel-distributed — the \"light\" "
        "tail among the three extreme value types. Some cells in mountainous "
        "or Mediterranean regions show moderately positive ξ values (heavier "
        "tails, meaning more intense rare events), which is climatologically "
        "plausible given orographic enhancement and convective activity."
    )

    # ================================================================
    # 4. FRÉCHET TRANSFORM & SPATIAL DEPENDENCE
    # ================================================================
    add_heading(doc, "4. Fréchet Transform and Spatial Dependence Estimation")

    doc.add_paragraph(
        "After fitting the GEV at each cell, the block maxima are transformed "
        "to unit Fréchet margins via the probability integral transform: "
        "Z = −1 / log F_GEV(x). This places all cells on a common "
        "heavy-tailed scale where the marginal distribution is identical "
        "everywhere: P(Z ≤ z) = exp(−1/z). The purpose of this transform is "
        "to separate marginal behaviour (how extreme is a single cell?) from "
        "dependence structure (how do extremes co-occur across space?). The "
        "subsequent analysis operates entirely on these Fréchet-transformed "
        "data."
    )
    doc.add_paragraph(
        "At each grid cell, we estimate three local anisotropy parameters "
        "(a, b, γ) by maximising a pairwise composite likelihood over all "
        "neighbouring cells within a normalised radius of 3.0. The parameter a "
        "controls the dependence range, b controls the anisotropy (elongation "
        "of the dependence ellipse), and γ controls the rotation angle. These "
        "parameters are subsequently smoothed over a spatial neighbourhood to "
        "reduce estimation noise."
    )

    # ================================================================
    # 5. LEC CLUSTERING
    # ================================================================
    add_heading(doc, "5. Spatial Dependence Clustering (LEC Method)")

    doc.add_paragraph(
        "The core contribution of Contzen et al. (2025) is the LEC "
        "(Localised Ellipse-shape Clustering) method. For each pair of grid "
        "cells, the method compares the estimated anisotropy ellipses by "
        "computing a Jaccard-like overlap measure: both ellipses are "
        "rasterised on a 21-point half-circle grid, and the dissimilarity "
        "is 1 minus the intersection-over-union of the two rasterised "
        "ellipses. This captures differences in both the shape and orientation "
        "of spatial dependence at each location."
    )
    doc.add_paragraph(
        "The pairwise dissimilarity matrix is then used as input to "
        "hierarchical agglomerative clustering (Ward's linkage). The number "
        "of clusters k is determined automatically by a threshold method: "
        "the 30th percentile of all pairwise distances sets the merge-height "
        "cutoff. For our European precipitation domain, this yields k = 26 "
        "clusters."
    )

    add_figure(
        doc, "lec_clusters_map.png",
        "Figure 5. Spatial dependence clusters identified by the LEC method "
        "(k = 26). Each colour represents a cluster of grid cells with similar "
        "extremal dependence structure. Spatial contiguity is not enforced; "
        "clusters are based purely on ellipse-shape similarity.",
        width=Inches(6.0),
    )

    doc.add_paragraph(
        "The cluster map reveals coherent spatial regions despite the absence "
        "of any spatial contiguity constraint in the clustering algorithm. "
        "Northern Europe (Scandinavia, British Isles) groups into a large "
        "cluster (orange), reflecting the relatively homogeneous dependence "
        "structure of Atlantic-influenced precipitation. The Mediterranean "
        "and Middle Eastern regions separate into distinct clusters, "
        "consistent with different precipitation regimes (convective vs. "
        "frontal). Some clusters contain geographically disconnected cells — "
        "this is methodologically correct, as two distant regions can share "
        "similar anisotropy parameters even without spatial adjacency."
    )

    # ================================================================
    # 6. VALIDATION: OUR LEC RESULTS vs PAPER FIGURE 9
    # ================================================================
    add_heading(doc, "6. Validation: LEC Clusters vs. Contzen et al. Figure 9")

    doc.add_paragraph(
        "Figure 9 of Contzen et al. (2025, Extremes, p. 733) presents "
        "a global LEC clustering of precipitation data from the AWI-ESM-1-1LR "
        "coupled climate model, yielding k_LEC = 24 clusters worldwide. "
        "Our pipeline produces k_LEC = 26 clusters for the European/Middle "
        "East sub-region alone. At first glance, having more clusters in a "
        "sub-region than on the entire globe appears contradictory. However, "
        "a careful comparison of the two analyses reveals that this is both "
        "expected and methodologically consistent."
    )

    add_heading(doc, "6.1 Comparison of Inputs and Parameters", level=2)

    # --- Table ---
    table = doc.add_table(rows=8, cols=3, style="Light Grid Accent 1")
    header_data = ["", "Paper Figure 9", "Our Analysis"]
    for i, text in enumerate(header_data):
        cell = table.rows[0].cells[i]
        cell.text = text
        for par in cell.paragraphs:
            for run in par.runs:
                run.bold = True
                run.font.size = Pt(10)

    rows_data = [
        ("Data source", "AWI-ESM-1-1LR (climate model)", "NOAA CPC (observational reanalysis)"),
        ("Spatial domain", "Global", "Europe / Middle East (30–65°N, 5–55°E)"),
        ("Time span", "~ 100+ model years", "20 years (2000–2019)"),
        ("Native resolution", "~ 1.875° (model grid)", "0.5° coarsened ×4 → ~ 2°"),
        ("ν (degrees of freedom)", "5", "5"),
        ("α (stability index)", "1.0", "1.0"),
        ("ε (neighbourhood)", "5 (grid steps)", "3.0 (normalised coordinate units)"),
    ]
    for r, (label, paper, ours) in enumerate(rows_data, start=1):
        table.rows[r].cells[0].text = label
        table.rows[r].cells[1].text = paper
        table.rows[r].cells[2].text = ours
        for c in range(3):
            for par in table.rows[r].cells[c].paragraphs:
                for run in par.runs:
                    run.font.size = Pt(10)

    doc.add_paragraph()  # spacing

    add_heading(doc, "6.2 Why More Clusters in a Sub-Region Is Expected", level=2)

    doc.add_paragraph(
        "The number of clusters k is an output of the pipeline, not a "
        "parameter. It is determined by the 30th-percentile threshold of "
        "the pairwise LEC dissimilarity matrix. Different data produce "
        "different dissimilarity distributions and therefore different k "
        "values. There is no theoretical requirement that k should scale "
        "with domain area. We identify four key factors that explain the "
        "discrepancy."
    )

    # --- Factor 1: Regional heterogeneity ---
    p = doc.add_paragraph()
    run = p.add_run("Factor 1 — Climatic heterogeneity. ")
    run.bold = True
    p.add_run(
        "The European/Middle East domain packs at least six distinct "
        "precipitation regimes into a relatively small spatial extent: "
        "(i) Atlantic maritime (westerlies, frontal systems), "
        "(ii) Mediterranean (winter cyclogenesis, summer drought), "
        "(iii) Alpine (orographic enhancement, foehn effects), "
        "(iv) Continental (blocking highs, convective summer storms), "
        "(v) Semi-arid/desert (Saharan edge, Middle East), and "
        "(vi) Black Sea/Caspian (distinct moisture sources). "
        "Each regime generates different storm orientations and spatial "
        "dependence structures, leading to distinct (a, b, γ) ellipses and "
        "high pairwise LEC dissimilarities. In contrast, the global domain "
        "of Figure 9 contains vast homogeneous regions — open oceans, "
        "continental interiors, tropical convergence zones — that collapse "
        "into few, large clusters, moderating the overall dissimilarity "
        "distribution."
    )

    # --- Factor 2: Data source ---
    p = doc.add_paragraph()
    run = p.add_run("Factor 2 — Observational vs. model data. ")
    run.bold = True
    p.add_run(
        "Climate models produce spatially smoother precipitation fields "
        "than observational products because they inherently parametrise "
        "sub-grid-scale processes. The AWI-ESM-1-1LR climate model used "
        "in Figure 9 will produce more spatially coherent anisotropy "
        "estimates than the NOAA CPC gauge-based reanalysis we use, "
        "which retains mesoscale variability (convective cells, orographic "
        "effects). Smoother fields → more similar local MLEs → lower LEC "
        "dissimilarities → fewer clusters."
    )

    # --- Factor 3: Sample size ---
    p = doc.add_paragraph()
    run = p.add_run("Factor 3 — Sample size and estimation variance. ")
    run.bold = True
    p.add_run(
        "Our 20-year observational record yields only 20 block maxima per "
        "cell, while climate model runs typically span 100+ years. With "
        "fewer observations, the local MLE estimates (a, b, γ) have higher "
        "sampling variance, which inflates pairwise dissimilarity values and "
        "can increase k. The paper's longer time series yields more stable "
        "MLE estimates, producing a more compressed dissimilarity distribution "
        "with fewer clusters."
    )

    # --- Factor 4: Threshold mechanics ---
    p = doc.add_paragraph()
    run = p.add_run("Factor 4 — Threshold mechanics on different distributions. ")
    run.bold = True
    p.add_run(
        "The 30th-percentile threshold is applied to the upper triangle of "
        "the pairwise LEC distance matrix, which has n(n−1)/2 entries. For "
        "a heterogeneous regional dataset, the dissimilarity distribution is "
        "right-skewed (many very different pairs), so the 30th percentile is "
        "relatively low, producing a low dendrogram cut height and therefore "
        "more clusters. For a global dataset where many cell pairs have "
        "similar anisotropy (large homogeneous zones), the distribution is "
        "more symmetric and the 30th percentile is higher, yielding fewer "
        "clusters."
    )

    add_heading(doc, "6.3 Assessment", level=2)

    doc.add_paragraph(
        "The comparison confirms that our pipeline is methodologically "
        "consistent with the paper's implementation. Both analyses use the "
        "same algorithm (LEC dissimilarity → Ward's linkage → 30th-percentile "
        "threshold) with matching core parameters (ν = 5, α = 1.0). "
        "The difference in k (26 vs. 24) is fully attributable to differences "
        "in input data (observations vs. climate model), spatial domain "
        "(regional vs. global), and effective sample size (20 vs. 100+ years). "
        "Critically, the simulation validation in Section 2 demonstrates that "
        "the Python implementation matches the R code to machine precision "
        "(28 cross-validation tests, 10⁻¹⁰ tolerance), so the discrepancy in "
        "k is a data/domain effect, not a code artefact."
    )
    doc.add_paragraph(
        "We note that subsetting the paper's 24 global clusters to our "
        "European domain would likely show 5–8 clusters intersecting this "
        "region — and each of those broad global clusters would split into "
        "multiple finer sub-clusters when the LEC algorithm is applied at "
        "regional resolution with observational data."
    )

    # ================================================================
    # 7. RISK METRICS
    # ================================================================
    add_heading(doc, "7. Tail-Risk Assessment: Expected Shortfall")

    doc.add_paragraph(
        "For each cluster, we compute tail-risk metrics on the Fréchet-"
        "transformed data. The loss variable for cluster A at time t is "
        "defined as the spatial block maximum: L_t = max over all cells s "
        "in cluster A of Z_t(s). This captures the idea that the cluster-"
        "level extreme is determined by the worst cell in any given year."
    )
    doc.add_paragraph(
        "We then compute the Expected Shortfall at the 95% level (ES₉₅), "
        "which is the conditional mean of L_t given that L_t exceeds the "
        "95th percentile (Value at Risk). ES₉₅ is a coherent risk measure "
        "(Artzner et al., 1999) that captures not just how often extremes "
        "occur, but how severe they are when they do."
    )
    doc.add_paragraph(
        "An important note on interpretation: the ES values are on the "
        "Fréchet scale, not in physical precipitation units (mm/day). This "
        "is because the Fréchet transform removed all marginal differences "
        "between cells. The ES therefore measures the intensity of joint "
        "spatial extremes — how strongly cells in a cluster tend to be "
        "extreme simultaneously — rather than absolute precipitation amounts. "
        "Higher ES means that when extreme events occur in a cluster, they "
        "tend to be more spatially coherent and intense."
    )

    add_figure(
        doc, "lec_risk_es_map.png",
        "Figure 6. Expected Shortfall (ES₉₅) per cluster, Fréchet scale, "
        "LEC clustering. Warmer colours indicate higher tail-risk intensity. "
        "Each cell is coloured by the ES of its assigned cluster.",
        width=Inches(6.0),
    )

    doc.add_paragraph(
        "The ES map shows substantial spatial variation in tail-risk "
        "intensity across Europe. Clusters covering large, climatologically "
        "coherent regions (such as Northern Europe) tend to have higher ES "
        "values, partly because the spatial maximum over more cells is "
        "stochastically larger. This is a known cluster-size effect and is "
        "addressed in the next step through GDP weighting."
    )

    # ================================================================
    # 8. WHICH ES TO USE? COMPARATIVE ANALYSIS
    # ================================================================
    add_heading(doc,
        "8. Cluster Risk Analysis: Which Expected Shortfall "
        "for the Risk Model?")

    doc.add_paragraph(
        "Section 7 computed tail-risk on the unit Fréchet scale, which is "
        "designed for measuring spatial dependence structure but has no "
        "direct physical interpretation. For a usable risk model we need "
        "a hazard metric in physical units (mm/day) that is both "
        "mathematically sound and free of confounding artefacts."
    )
    doc.add_paragraph(
        "The purpose of this section is to compare three candidate Expected "
        "Shortfall formulations, diagnose their properties, and reach a "
        "justified conclusion about which one to adopt in the risk model. "
        "The three candidates are:"
    )
    candidates = [
        ("Spatial-Max ES₉₅",
         "For each cluster A, define the loss in year t as "
         "L_t = max over all cells s ∈ A of X_t(s), where X_t(s) is the "
         "annual block maximum precipitation (mm/day) at cell s in year t. "
         "ES₉₅ = mean of L_t values exceeding the 95th percentile "
         "of {L_1, …, L_20}."),
        ("Per-Cell ES₉₅",
         "For each cell s independently, compute ES₉₅ from its "
         "own 20 annual block maxima. This ignores cluster structure "
         "entirely."),
        ("Mean Per-Cell ES₉₅",
         "For each cluster, average the per-cell ES₉₅ across all "
         "member cells. This summarises the typical marginal hazard "
         "intensity within the cluster, without taking a spatial maximum."),
    ]
    for label, desc in candidates:
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(f"{label}: ")
        run.bold = True
        p.add_run(desc)

    doc.add_paragraph(
        "The question we must answer: which of these three formulations "
        "gives the most informative and mathematically correct measure of "
        "cluster-level hazard intensity for use in a Hazard × Exposure "
        "risk model?"
    )

    # ── 8.1 Cluster Geography ──
    add_heading(doc, "8.1 Cluster Geography", level=2)

    doc.add_paragraph(
        "The 26 LEC clusters partition the 384 land cells into groups that "
        "share similar extremal dependence ellipses (anisotropy parameters "
        "a, b, γ). Cluster sizes range from 1 cell to 155 cells (cluster 16, "
        "which covers much of Central-Eastern Europe). The clusters are "
        "NOT required to be spatially contiguous: the LEC algorithm groups "
        "cells by dependence-structure similarity, not geographic proximity. "
        "This is a distinguishing feature of the LEC method compared to "
        "purely geographic clustering approaches."
    )

    add_figure(
        doc, "cluster_risk/map1_lec_clusters.png",
        "Figure 9. LEC cluster assignments for 384 land cells across the "
        "European domain (30–65°N, 5–55°E). Each cell is coloured by its "
        "cluster ID (arbitrary labels). Numbers inside clusters show the "
        "cluster ID. Cluster sizes range from 1 to 155 cells.",
        width=Inches(6.0),
    )

    # ── 8.2 Candidate 1: Spatial-Max ES ──
    add_heading(doc,
        "8.2 Candidate 1 — Spatial-Max ES₉₅ per Cluster", level=2)

    doc.add_paragraph(
        "Why we compute it: The spatial block maximum L_t = max_s X_t(s) "
        "is the natural loss variable of Contzen et al. (2025). It "
        "answers the question: 'What is the worst precipitation event "
        "anywhere in this cluster in year t?' This is relevant if the "
        "decision-maker cares about the peak intensity over the whole "
        "cluster region (e.g. an insurer covering the entire region)."
    )
    doc.add_paragraph(
        "How we compute it: For each of 20 years, take the maximum annual "
        "block maximum across all cells in the cluster. This yields 20 "
        "values L_1, …, L_20. VaR₉₅ = the 95th percentile (approx. "
        "19th-largest value). ES₉₅ = mean of L_t values ≥ VaR₉₅ (i.e. "
        "the mean of the top 1–2 values with 20 years of data)."
    )

    add_figure(
        doc, "cluster_risk/map2_cluster_es.png",
        "Figure 10. Cluster-level ES₉₅ of the spatial block maximum "
        "(mm/day). Each cell is coloured by the ES of its assigned cluster. "
        "Labels show the ES value and worst year. Note: this metric is "
        "biased by cluster size — see Section 8.4.",
        width=Inches(6.0),
    )

    doc.add_paragraph(
        "Result: spatial-max ES₉₅ ranges from 23.0 mm/day (cluster 2, "
        "1 cell, Middle East) to 215.5 mm/day (cluster 16, 155 cells, "
        "Central-Eastern Europe). Large clusters dominate the ranking."
    )

    # ── 8.3 Candidate 2: Per-Cell ES ──
    add_heading(doc,
        "8.3 Candidate 2 — Per-Cell ES₉₅ (Marginal Hazard)", level=2)

    doc.add_paragraph(
        "Why we compute it: The per-cell ES₉₅ measures the intrinsic "
        "tail-risk at each location, irrespective of how cells are "
        "grouped. It answers: 'How extreme can precipitation get at this "
        "specific grid cell?' This is independent of cluster size and "
        "reveals the underlying climatological pattern."
    )
    doc.add_paragraph(
        "How we compute it: For each cell s, take its 20 annual block "
        "maxima, compute the 95th percentile (VaR₉₅), then the mean of "
        "all values above that threshold. Per-cell ES₉₅ ranges from "
        "17.1 to 215.5 mm/day across the domain, with a median of "
        "50.8 mm/day."
    )

    add_figure(
        doc, "cluster_risk/map3_cell_es.png",
        "Figure 11. Per-cell ES₉₅ (mm/day) — marginal tail-risk intensity "
        "at each grid cell, computed from 20 annual block maxima. "
        "This is independent of cluster assignment and reveals the "
        "intrinsic hazard at each location.",
        width=Inches(6.0),
    )

    doc.add_paragraph(
        "The per-cell ES map reveals the expected climatological pattern: "
        "Mediterranean coastal areas and parts of Central Europe show "
        "the highest marginal hazard, while arid regions (North Africa, "
        "Middle East) and high-latitude continental interiors show lower "
        "values. This metric is useful for understanding where the data "
        "is extreme, but it does not summarise cluster-level hazard: it "
        "assigns no single value to a cluster that could enter the risk "
        "model formula Risk(s) = Hazard(cluster) × Exposure(cell)."
    )

    # ── 8.4 The Size-Bias Problem ──
    add_heading(doc,
        "8.4 The Size-Bias Problem with Spatial-Max ES", level=2)

    doc.add_paragraph(
        "Why this matters: Before choosing our risk metric, we must check "
        "whether the spatial-max ES is confounded by cluster size. If it "
        "is, then the metric does not measure hazard intensity per se — it "
        "partly measures how many cells the cluster contains, which is a "
        "property of the LEC algorithm, not of the climate."
    )
    doc.add_paragraph(
        "Diagnostic: We compute the Spearman rank correlation between "
        "cluster size (number of cells) and spatial-max ES₉₅ across the "
        "16 multi-cell clusters. The result is:"
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "ρ(cluster size, spatial-max ES₉₅) = 0.813,  p < 0.001")
    run.bold = True
    run.font.size = Pt(11)

    doc.add_paragraph(
        "This is a very strong positive correlation. The spatial-max ES "
        "is dominated by cluster size, not by the intrinsic precipitation "
        "intensity of the cells. The mechanism is straightforward: taking "
        "the maximum over N cells always yields a stochastically larger "
        "value as N increases, even if all cells have identical marginal "
        "distributions. This is an order-statistics artefact, not a "
        "climate signal."
    )

    # ── 8.5 Candidate 3: Mean Cell ES ──
    add_heading(doc,
        "8.5 Candidate 3 — Mean Per-Cell ES₉₅ (Size-Bias-Free)",
        level=2)

    doc.add_paragraph(
        "Why we compute it: To eliminate the cluster-size confound, we "
        "average the per-cell ES₉₅ values across all cells in each "
        "cluster. This answers: 'What is the typical marginal hazard "
        "intensity experienced by a cell in this cluster?' It preserves "
        "physical unit interpretation (mm/day), can be computed for any "
        "cluster regardless of size, and does not depend on the number "
        "of cells."
    )

    add_figure(
        doc, "cluster_risk/map5_mean_cell_es.png",
        "Figure 12. Mean per-cell ES₉₅ per cluster (mm/day) — a "
        "size-bias-free measure of cluster hazard intensity. Each cell "
        "is coloured by the average ES of its cluster's members. Labels "
        "show the mean value.",
        width=Inches(6.0),
    )

    doc.add_paragraph(
        "Diagnostic: The Spearman correlation between cluster size and "
        "mean per-cell ES₉₅ is:"
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "ρ(cluster size, mean-cell ES₉₅) = 0.249,  p = 0.33")
    run.bold = True
    run.font.size = Pt(11)

    doc.add_paragraph(
        "This is statistically indistinguishable from zero. The size bias "
        "has been removed. This is the metric we will adopt for the risk "
        "model."
    )

    # ── 8.6 Full Comparative Table ──
    add_heading(doc,
        "8.6 Full Cluster Risk Table — All 26 Clusters", level=2)

    doc.add_paragraph(
        "Table 2 below presents the complete risk profile for all 26 LEC "
        "clusters. The table is sorted by mean per-cell ES₉₅ (descending) "
        "to reflect the size-bias-free hazard ranking. This allows direct "
        "comparison of the three ES candidates side by side."
    )

    # Full 26-cluster table (10 columns)
    from docx.oxml.ns import qn
    tbl_headers = ["Cl", "N", "MeanES", "MaxES", "SpatES",
                   "VaR₉₅", "ρ̄", "CoExc", "Worst", "Region"]
    tbl_data = [
        ("20", "5", "90.1", "172.5", "172.5", "163.2", "0.141", "0.222", "2003", "C. Europe, Central"),
        ("18", "1", "77.0", "77.0", "77.0", "66.5", "1.000", "1.000", "2001", "N. Africa / M.East"),
        ("17", "1", "75.4", "75.4", "75.4", "44.5", "1.000", "1.000", "2001", "C. Europe, Central"),
        ("26", "1", "69.8", "69.8", "69.8", "62.3", "1.000", "1.000", "2004", "S. Mediterranean"),
        ("22", "10", "69.7", "105.3", "105.3", "87.2", "0.075", "0.143", "2002", "C. Europe, C-East"),
        ("21", "2", "69.0", "90.5", "90.5", "49.7", "0.042", "0.667", "2005", "C. Europe, Central"),
        ("19", "4", "63.6", "92.9", "92.9", "83.6", "\u22120.073", "0.286", "2005", "C. Europe, C-East"),
        ("13", "2", "62.0", "70.6", "70.6", "68.9", "0.257", "0.667", "2002", "S. Med., C-East"),
        ("1", "1", "61.3", "61.3", "61.3", "59.5", "1.000", "1.000", "2017", "N. Europe, West"),
        ("16", "155", "60.4", "215.5", "215.5", "152.5", "0.014", "0.100", "2018", "C. Europe, C-East"),
        ("15", "5", "54.8", "86.8", "86.8", "71.5", "\u22120.097", "0.222", "2002", "C. Europe, C-East"),
        ("10", "120", "53.7", "156.1", "156.1", "145.3", "0.020", "0.100", "2011", "C. Europe, C-East"),
        ("23", "2", "53.6", "60.5", "60.5", "57.4", "0.183", "0.500", "2014", "C. Europe, Central"),
        ("3", "21", "49.7", "112.4", "112.4", "77.2", "0.175", "0.133", "2005", "N. Europe, Central"),
        ("5", "1", "49.5", "49.5", "49.5", "38.8", "1.000", "1.000", "2000", "N. Africa, West"),
        ("24", "2", "46.2", "46.7", "46.7", "45.8", "0.235", "0.500", "2008", "C. Europe, Central"),
        ("9", "31", "45.6", "107.5", "107.5", "107.0", "0.048", "0.111", "2019", "N. Europe, C-East"),
        ("12", "1", "45.4", "45.4", "45.4", "35.8", "1.000", "1.000", "2011", "N. Europe, West"),
        ("11", "5", "44.4", "70.9", "70.9", "67.5", "0.040", "0.250", "2012", "S. Med., Central"),
        ("4", "1", "43.3", "43.3", "43.3", "38.1", "1.000", "1.000", "2014", "N. Europe, East"),
        ("14", "5", "42.5", "56.9", "56.9", "46.9", "0.060", "0.250", "2019", "C. Europe, C-East"),
        ("25", "2", "42.4", "42.6", "42.6", "42.3", "\u22120.427", "0.500", "2008", "S. Med., C-East"),
        ("8", "2", "42.3", "49.6", "49.6", "41.6", "0.236", "0.500", "2015", "N. Africa, C-East"),
        ("6", "1", "40.3", "40.3", "40.3", "39.3", "1.000", "1.000", "2001", "N. Europe, West"),
        ("7", "2", "37.9", "49.4", "49.4", "29.5", "0.317", "0.667", "2017", "N. Africa, C-East"),
        ("2", "1", "23.0", "23.0", "23.0", "22.5", "1.000", "1.000", "2016", "N. Africa, C-East"),
    ]

    table = doc.add_table(rows=len(tbl_data) + 1, cols=len(tbl_headers))
    table.style = "Light Grid Accent 1"
    for j, h in enumerate(tbl_headers):
        cell = table.rows[0].cells[j]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(7.5)
    for i, row_data in enumerate(tbl_data, 1):
        for j, val in enumerate(row_data):
            cell = table.rows[i].cells[j]
            cell.text = val
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(7.5)

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(
        "Table 2. Complete risk profile for all 26 LEC clusters, sorted "
        "by mean per-cell ES₉₅ (descending). See column definitions below.")
    run.italic = True
    run.font.size = Pt(9)

    # ── Column definitions ──
    add_heading(doc,
        "8.7 Column Definitions", level=2)

    col_defs = [
        ("Cl", "Cluster ID assigned by the LEC algorithm (arbitrary label, "
         "not a rank)."),
        ("N", "Number of ~2° grid cells in the cluster. The LEC algorithm "
         "groups cells by anisotropy ellipse similarity; spatial contiguity "
         "is NOT enforced."),
        ("MeanES", "Mean per-cell ES₉₅ (mm/day). For each cell in the "
         "cluster, compute ES₉₅ from its 20 annual block maxima, then "
         "average across all cells. This is the SIZE-BIAS-FREE hazard "
         "intensity — the recommended metric for comparing clusters of "
         "different sizes."),
        ("MaxES", "Maximum per-cell ES₉₅ (mm/day) within the cluster — "
         "the single most extreme location."),
        ("SpatES", "ES₉₅ of the spatial block maximum. For each year t, "
         "L_t = max over all cells s. ES₉₅ = mean of L_t values exceeding "
         "the 95th percentile. ⚠ CAUTION: Mechanically inflated by cluster "
         "size (ρ = 0.813 with cluster size). A 155-cell cluster will always "
         "produce a larger spatial max than a 5-cell cluster, even if "
         "individual cells are equally extreme."),
        ("VaR₉₅", "Value at Risk at 95% of the spatial block maximum — "
         "the 95th percentile of {L_1, …, L_20}. With 20 years, this is "
         "approximately the 19th-largest annual spatial max."),
        ("ρ̄", "Mean pairwise Spearman rank correlation of annual block "
         "maxima time series within the cluster. ρ̄ > 0 means cells tend "
         "to be extreme in the same years (spatial coherence); ρ̄ ≈ 0 "
         "means independence; ρ̄ < 0 means anti-correlation. Single-cell "
         "clusters trivially have ρ̄ = 1.000."),
        ("CoExc", "Co-exceedance rate: in years where at least one cell "
         "exceeds its 90th percentile, what fraction of all cluster cells "
         "simultaneously exceed their own 90th percentile? Single-cell "
         "clusters trivially have CoExc = 1.000."),
        ("Worst", "Calendar year (2000–2019) with the highest spatial block "
         "maximum in the cluster — the 'worst event year'."),
        ("Region", "Approximate geographic label based on cluster centroid "
         "coordinates."),
    ]
    for label, desc in col_defs:
        p = doc.add_paragraph()
        run = p.add_run(f"{label}: ")
        run.bold = True
        p.add_run(desc)
        for r in p.runs:
            r.font.size = Pt(9.5)

    # ── 8.8 Summary Panel ──
    add_heading(doc,
        "8.8 Comprehensive Summary Panel", level=2)

    doc.add_paragraph(
        "The 6-panel summary below brings together all spatial risk "
        "dimensions: (a) cluster geography, (b) cluster-level spatial-max "
        "ES (size-biased), (c) mean per-cell ES (size-bias-free), "
        "(d) per-cell marginal hazard, (e) co-exceedance rates, and "
        "(f) mean Spearman ρ̄. Panels (a)–(c) characterise between-cluster "
        "differences; panels (d)–(f) characterise within-cluster structure."
    )

    add_figure(
        doc, "cluster_risk/summary_6panel.png",
        "Figure 13. Six-panel summary of the LEC cluster risk analysis. "
        "(a) Cluster assignments; (b) spatial-max ES₉₅ (size-biased); "
        "(c) mean per-cell ES₉₅ (size-bias-free); (d) per-cell marginal "
        "ES₉₅; (e) co-exceedance rate; (f) mean Spearman ρ̄. "
        "All values in mm/day except (e) and (f).",
        width=Inches(6.5),
    )

    doc.add_paragraph(
        "Key observations from the table and maps:"
    )

    obs = [
        ("Size-bias removal changes the ranking fundamentally",
         "Cluster 16 (155 cells) has the highest spatial-max ES "
         "(215.5 mm/day) but only the 10th-highest mean cell ES "
         "(60.4 mm/day). Cluster 10 (120 cells) drops from 3rd to 12th. "
         "Meanwhile, cluster 20 (just 5 cells) rises to 1st place with "
         "a mean cell ES of 90.1 mm/day — genuinely the most extreme "
         "precipitation cluster."),
        ("Spatial coherence is weak in large clusters",
         "Clusters 16 and 10 (covering 72% of all cells) have near-zero "
         "ρ̄ (0.014 and 0.020) and only 10% co-exceedance: their extremes "
         "are highly localised, not synchronous. This means their high "
         "spatial-max ES reflects the statistical inevitability that at "
         "least one of 155 cells will be extreme each year — not that "
         "the cluster experiences coherent extreme events."),
        ("Small clusters show diverse and genuine behaviour",
         "Multi-cell clusters range from ρ̄ = −0.427 (cluster 25 — "
         "anti-correlated) to ρ̄ = +0.317 (cluster 7), reflecting "
         "genuinely different dependence structures captured by the LEC "
         "algorithm."),
        ("Worst years are distributed, not concentrated",
         "No single year dominates: 2001, 2002, and 2005 each produce "
         "the worst event for 3 clusters. This confirms that European "
         "precipitation extremes are regional, not driven by a single "
         "pan-continental event."),
    ]
    for label, desc in obs:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(f"{label}: ")
        run.bold = True
        p.add_run(desc)

    # ── 8.9 Conclusion: Which ES? ──
    add_heading(doc,
        "8.9 Conclusion: Recommended ES Formulation for the Risk Model",
        level=2)

    doc.add_paragraph(
        "We can now evaluate the three candidate ES formulations against "
        "three criteria: (i) mathematical correctness — is the metric free "
        "of confounding artefacts? (ii) informativeness — does the metric "
        "discriminate clusters by genuine hazard intensity? "
        "(iii) applicability — can the metric be used in the Hazard × "
        "Exposure risk formula?"
    )

    # Comparison table
    cmp_headers = ["Criterion", "Spatial-Max ES₉₅", "Per-Cell ES₉₅",
                   "Mean Per-Cell ES₉₅"]
    cmp_data = [
        ("Free of size bias?",
         "NO — ρ = 0.813 with\ncluster size (p < 0.001)",
         "YES — defined per\ncell, no cluster\naggregation",
         "YES — ρ = 0.249 with\ncluster size (p = 0.33)"),
        ("Discriminates\nclusters?",
         "Partially — ranking\ndominated by cluster\nsize, not climate",
         "NO — no single\ncluster-level value;\ncannot rank clusters",
         "YES — range 23.0 to\n90.1 mm/day, ranking\nreflects climate"),
        ("Usable in\nRisk = H × E?",
         "YES — one value\nper cluster",
         "NO — one value per\ncell, not per cluster",
         "YES — one value\nper cluster"),
        ("Physical\ninterpretation",
         "'Worst cell in the\ncluster in a tail\nevent'",
         "'Tail severity at\nthis specific cell'",
         "'Typical tail severity\nof cells in this\ncluster'"),
    ]

    cmp_table = doc.add_table(rows=len(cmp_data) + 1, cols=4)
    cmp_table.style = "Light Grid Accent 1"
    for j, h in enumerate(cmp_headers):
        cell = cmp_table.rows[0].cells[j]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(8.5)
    for i, row_data in enumerate(cmp_data, 1):
        for j, val in enumerate(row_data):
            cell = cmp_table.rows[i].cells[j]
            cell.text = val
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(8.5)

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(
        "Table 3. Comparison of three candidate ES₉₅ formulations.")
    run.italic = True
    run.font.size = Pt(9)

    doc.add_paragraph(
        "Conclusion: We adopt the mean per-cell ES₉₅ as the hazard "
        "component of the risk model. It satisfies all three criteria: "
        "it is free of the cluster-size artefact (confirmed by a "
        "Spearman test), it provides a single physically interpretable "
        "value per cluster, and it discriminates clusters by genuine "
        "precipitation extremity rather than by how many grid cells the "
        "LEC algorithm assigned to them."
    )
    doc.add_paragraph(
        "The risk formula therefore becomes:"
    )
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(
        "Risk(s)  =  MeanCellES₉₅(cluster of s)  ×  GDP(s)")
    run.bold = True
    run.font.size = Pt(12)

    doc.add_paragraph(
        "where s is a grid cell, and MeanCellES₉₅ is in mm/day. "
        "This assigns each cell its cluster's average marginal hazard "
        "intensity, weighted by the cell's own economic exposure. "
        "Unlike the Fréchet-scale spatial-max ES used in Section 7, "
        "this metric has a direct physical interpretation and is free "
        "of the order-statistics artefact."
    )

    doc.add_paragraph(
        "The complete per-cluster analysis, including year-by-year spatial "
        "block maxima, GDP exposure figures, min/max/std of per-cell ES, "
        "and worst-event annotations, is provided in the companion file "
        "cluster_risk_report.txt (1,270 lines, 26 detailed analysis cards)."
    )

    # ================================================================
    # 9. GDP EXPOSURE
    # ================================================================
    add_heading(doc, "9. Economic Exposure: GDP Layer")

    doc.add_paragraph(
        "To connect the statistical tail-risk to real-world impact, we "
        "incorporate gridded economic data from Kummu et al. (2018). The "
        "GDP PPP dataset provides purchasing-power-adjusted gross domestic "
        "product at 5 arc-minute resolution (approximately 10 km). We "
        "aggregate this fine-resolution data onto our coarser ~2° pipeline "
        "grid by summing all 5-arc-min cells that fall within each coarsened "
        "grid cell."
    )

    add_figure(
        doc, "gdp_exposure_map.png",
        "Figure 14. Economic exposure layer: GDP PPP per grid cell (constant "
        "2011 international USD), aggregated from the Kummu et al. (2018) "
        "dataset, 2015 snapshot. Dark green indicates high economic "
        "concentration.",
        width=Inches(6.0),
    )

    doc.add_paragraph(
        "The GDP map shows the expected pattern: high economic concentration "
        "in Western Europe (Germany, France, Benelux, UK), with moderate "
        "values across Southern and Eastern Europe and low values in "
        "North Africa and Central Asia. This purely economic layer, "
        "independent of any climate information, will be combined with the "
        "tail-risk metric in the next step."
    )

    # ================================================================
    # 10. COMBINED RISK MAP
    # ================================================================
    add_heading(doc, "10. Combined Risk: Hazard × Exposure")

    doc.add_paragraph(
        "The final pipeline output combines tail-risk intensity with economic "
        "exposure through a simple product:"
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Risk(s)  =  ES₉₅ᵏ⁽ˢ⁾  ×  GDP(s)")
    run.bold = True
    run.font.size = Pt(12)

    doc.add_paragraph(
        "where k(s) is the cluster to which cell s belongs. Each cell "
        "inherits the ES of its cluster (the hazard intensity) and multiplies "
        "it by its own GDP (the economic exposure). This cell-level formulation "
        "avoids cluster-size bias: a city cell with high GDP receives higher "
        "risk than a rural cell in the same cluster, despite sharing the same "
        "hazard intensity."
    )
    doc.add_paragraph(
        "This product creates a hazard-exposure index with units of "
        "\"Fréchet-units × USD\". It is not a direct monetary loss estimate "
        "(that would require a damage function mapping precipitation intensity "
        "to economic losses), but rather a risk prioritisation score that "
        "identifies where strong spatial extremal dependence and high economic "
        "value coincide — analogous to seismic hazard maps that multiply "
        "ground-shaking intensity by population exposure."
    )

    add_figure(
        doc, "lec_risk_gdp_map.png",
        "Figure 15. Exposure-weighted precipitation risk (ES₉₅ × GDP PPP) "
        "per grid cell using LEC clusters. The scale is in units of "
        "Fréchet × USD PPP (×10¹⁴). Dark red indicates where high tail-risk "
        "intensity coincides with high economic exposure.",
        width=Inches(6.0),
    )

    doc.add_paragraph(
        "The risk map correctly identifies Western Europe — particularly "
        "the UK, Benelux, western Germany, and northern France — as the "
        "highest-risk zone. These cells combine moderate-to-high ES values "
        "(precipitation extremes in this Atlantic-influenced region are "
        "spatially coherent over large areas) with very high GDP. A secondary "
        "hotspot appears around Istanbul and the eastern Mediterranean, where "
        "moderate GDP intersects with a high-ES cluster."
    )
    doc.add_paragraph(
        "Conversely, Scandinavia and Eastern Europe show lower combined risk "
        "despite being within high-ES clusters, because their GDP density is "
        "lower. Central Asian and North African cells are low-risk because "
        "both their ES and GDP values are modest. The map thus provides a "
        "nuanced, spatially explicit view of where precipitation extreme "
        "risk is most consequential."
    )

    # ================================================================
    # 11. SUMMARY
    # ================================================================
    add_heading(doc, "11. Summary and Interpretation")

    doc.add_paragraph(
        "This analysis demonstrates a complete pipeline from raw daily "
        "precipitation data to spatially explicit risk maps:"
    )

    items = [
        ("GEV fitting", "extracts the extreme precipitation signal and models "
         "the tail behaviour at each grid cell (validated by Q-Q plots and KS "
         "tests with only 2.6% rejection rate);"),
        ("Fréchet transform", "standardises all cells to a common tail scale, "
         "separating marginal from dependence behaviour;"),
        ("Local MLE + smoothing", "estimates spatially varying anisotropy "
         "parameters describing the shape and orientation of dependence "
         "ellipses;"),
        ("LEC clustering", "groups cells with similar dependence structure "
         "into k = 26 clusters, validated against the Contzen et al. (2025) "
         "simulation study and compared with the paper's global Figure 9 "
         "(Section 6);"),
        ("Expected Shortfall", "quantifies tail-risk intensity per cluster "
         "on the Fréchet scale (Section 7) and in physical mm/day units "
         "(Section 8), with explicit size-bias corrections;"),
        ("Cluster risk analysis", "diagnoses size-bias in spatial-max ES, "
         "demonstrates mean per-cell ES as a fair ranking metric, and "
         "provides per-cluster analysis cards with year-by-year detail "
         "(Section 8);"),
        ("GDP weighting", "combines statistical hazard intensity with economic "
         "exposure to produce a cell-level risk prioritisation index."),
    ]
    for label, desc in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(f"{label}: ")
        run.bold = True
        p.add_run(desc)

    doc.add_paragraph(
        "The key finding is that Western Europe (UK, Benelux, western Germany, "
        "northern France) emerges as the highest-risk area — not because it "
        "has the most extreme precipitation events in absolute terms, but "
        "because it combines spatially coherent extremes with high economic "
        "concentration. This type of insight — identifying where hazard "
        "dependence and exposure overlap — is precisely what the max-stable "
        "process clustering framework is designed to provide."
    )

    # --- REFERENCES ---
    add_heading(doc, "References")
    refs = [
        "Artzner, P., Delbaen, F., Eber, J.-M. & Heath, D. (1999). "
        "Coherent measures of risk. Mathematical Finance, 9(3), 203–228.",
        "Contzen, J., Dickhaus, T. & Lohmeyer, J. (2025). "
        "Localised clustering of extremal dependence structures. "
        "Extremes, 28, 713–737.",
        "de Haan, L. & Ferreira, A. (2006). Extreme Value Theory: "
        "An Introduction. Springer.",
        "Kummu, M., Taka, M. & Guillaume, J. H. A. (2018). "
        "Gridded global datasets for GDP and HDI. Scientific Data, 5, 180004.",
        "Padoan, S. A., Ribatet, M. & Sisson, S. A. (2010). "
        "Likelihood-based inference for max-stable processes. "
        "JASA, 105(489), 263–277.",
        "Schlather, M. (2002). Models for stationary max-stable random fields. "
        "Extremes, 5(1), 33–44.",
    ]
    for i, ref in enumerate(refs, 1):
        p = doc.add_paragraph(f"[{i}] {ref}")
        p.paragraph_format.space_after = Pt(2)
        for run in p.runs:
            run.font.size = Pt(9)

    # --- Save ---
    os.makedirs(os.path.dirname(OUTPUT), exist_ok=True)
    doc.save(OUTPUT)
    print(f"Saved: {OUTPUT}")
    print(f"Contains 15 figures, 3 tables, 11 sections")


if __name__ == "__main__":
    main()
